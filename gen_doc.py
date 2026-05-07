from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estilos globales ──────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for s in ['Heading 1','Heading 2','Heading 3']:
    hs = doc.styles[s]
    hs.font.name = 'Calibri'
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

def h1(t): doc.add_heading(t, 1)
def h2(t): doc.add_heading(t, 2)
def h3(t): doc.add_heading(t, 3)
def p(t=''):  doc.add_paragraph(t)
def b(t): doc.add_paragraph(t, style='List Bullet')

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()

def endpoint_table(rows):
    table(['Método', 'Ruta (vía Gateway)', 'Acceso requerido', 'Descripción'], rows)

# ════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TechFix Software')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Documentación Técnica del Sistema')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run('Versión 1.0  ·  Mayo 2026')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. QUÉ ES TECHFIX
# ════════════════════════════════════════════════════════════
h1('1. ¿Qué es TechFix?')
p('TechFix es un sistema de gestión para un servicio técnico de reparación. Permite a los clientes agendar visitas, consultar el catálogo de repuestos disponibles y ver el estado de sus órdenes. Al mismo tiempo, el personal puede administrar citas, registrar cobros y mantener el inventario actualizado, todo desde una sola plataforma web.')
p('El sistema está diseñado para funcionar de forma completa con un solo comando de despliegue. No requiere configuraciones manuales: al levantarlo, las bases de datos y los servicios quedan listos para usarse de inmediato.')

# ════════════════════════════════════════════════════════════
# 2. MÓDULOS
# ════════════════════════════════════════════════════════════
h1('2. Módulos del sistema')
p('El sistema está dividido en cinco áreas funcionales independientes. Si una falla, las demás siguen operando sin verse afectadas.')

h2('2.1 Usuarios y Roles')
p('Gestiona todas las cuentas. Existen tres tipos de usuario:')
b('USER — Cliente. Puede agendar citas, ver repuestos y consultar sus órdenes.')
b('TECNICO — Personal técnico. Puede además modificar el estado de citas y órdenes.')
b('ADMIN — Administrador único. Tiene acceso total: gestiona pagos, inventario, usuarios y asigna roles.')
p('Solo puede existir un administrador a la vez; el sistema lo controla automáticamente y no permite eliminarlo.')

h2('2.2 Pagos')
p('Registra y administra los cobros asociados a cada visita técnica. Cada pago guarda:')
b('El monto cobrado')
b('El método de pago: EFECTIVO, TARJETA_DEBITO, TARJETA_CREDITO, TRANSFERENCIA')
b('El estado: PENDIENTE, PAGADO, ANULADO')
b('Una descripción opcional y la fecha del cobro')
p('Incluye un mecanismo de protección contra cobros duplicados (idempotencia). Si por error se intenta registrar el mismo pago dos veces (por ejemplo, un doble clic), el sistema detecta la duplicación y devuelve el registro original sin crear un nuevo cobro. Esto se logra mediante una clave única (UUID) generada en el navegador y enviada junto a cada solicitud de pago.')

h2('2.3 Inventario (Stock / Repuestos)')
p('Administra el catálogo de productos disponibles para la venta o uso en reparaciones. Permite:')
b('Crear y organizar productos por categorías')
b('Consultar stock con filtros (categoría, estado, búsqueda por texto, SKU)')
b('Registrar movimientos de inventario: PURCHASE, SALE, ADJUSTMENT, RETURN, TRANSFER, DAMAGE')
b('Ajustar el stock directamente y ver el historial de movimientos')
b('Ver alertas de productos con stock bajo o agotado')
p('El catálogo es visible para cualquier visitante sin necesidad de iniciar sesión. Solo el administrador puede modificarlo.')

h2('2.4 Agenda de Citas')
p('Permite agendar y gestionar visitas técnicas. Cada cita registra la fecha y hora, el tipo de servicio (REPARACION o INSTALACION), el cliente, el técnico asignado y el estado actual:')
b('PENDIENTE — Cita recién creada')
b('EN_PROCESO — El técnico está atendiendo')
b('COMPLETADA — Visita finalizada')
b('CANCELADA — Cita cancelada')
p('El módulo también gestiona los registros de clientes y técnicos de forma independiente.')

h2('2.5 API Gateway')
p('Es el punto de entrada único del sistema. Toda solicitud del navegador pasa primero por aquí. El Gateway se encarga de verificar si el usuario está autenticado (comprobando su ID en el servicio de usuarios), validar que tiene los permisos necesarios para lo que intenta hacer, y redirigir la solicitud al servidor correcto.')
p('Rutas públicas que no requieren autenticación:')
b('Buscar usuario por nombre (login)')
b('Crear cuenta nueva (registro)')
b('Ver catálogo de repuestos (GET)')
b('Ver lista de técnicos (GET)')

# ════════════════════════════════════════════════════════════
# 3. ARQUITECTURA
# ════════════════════════════════════════════════════════════
h1('3. Arquitectura del sistema')
p('El navegador nunca habla directamente con los servidores individuales. Todo pasa por el Gateway (puerto 8090), que actúa como recepción centralizada:')
p('Navegador  →  Gateway (:8090)\n'
  '                ├─→  Usuarios  (:8080)\n'
  '                ├─→  Pagos     (:8081)\n'
  '                ├─→  Stock     (:8082)\n'
  '                └─→  Agenda    (:8083)')
p('Cada módulo tiene su propia base de datos PostgreSQL. No comparten datos entre sí, lo que evita que un problema en un área afecte a las demás.')
p('Para identificar al usuario, el frontend envía su ID en el header X-User-Id de cada petición. El Gateway consulta ese ID al servicio de usuarios para confirmar que existe, que está activo y qué rol tiene antes de dejar pasar la solicitud.')

table(
    ['Capa', 'Tecnología', 'Para qué sirve'],
    [
        ('Interfaz web', 'React + Vite', 'La página que ve el usuario'),
        ('Estilo visual', 'CSS puro', 'Diseño y apariencia'),
        ('Servidor web prod.', 'Nginx', 'Sirve la página al navegador'),
        ('Backend (5 servidores)', 'Java 21 + Spring Boot 3', 'Lógica del sistema'),
        ('Base de datos', 'PostgreSQL 16', 'Almacenamiento persistente'),
        ('Contenerización', 'Docker + Docker Compose', 'Levanta todo el stack junto'),
    ]
)

# ════════════════════════════════════════════════════════════
# 4. CÓMO LEVANTAR
# ════════════════════════════════════════════════════════════
h1('4. Cómo levantar el sistema')
p('Requisito previo: tener Docker Desktop instalado y corriendo.')
p('Pasos:')
b('Clonar el repositorio:  git clone <URL>  &&  cd TechFix')
b('Levantar todo:  docker compose up --build -d')
p('El primer arranque tarda varios minutos porque Docker descarga las imágenes base y compila los proyectos Java. Las siguientes veces es mucho más rápido.')

table(
    ['Servicio', 'URL local'],
    [
        ('Página web', 'http://localhost:3000'),
        ('API Gateway', 'http://localhost:8090'),
    ]
)

p('Para detener el sistema:  docker compose down')
p('Para detener y borrar también las bases de datos:  docker compose down -v')

h2('4.1 Orden de arranque')
p('Docker respeta un orden para que ningún servicio arranque antes de que su base de datos esté lista:')
b('1. Las cuatro bases de datos PostgreSQL')
b('2. Los cuatro servidores (roles, pagos, stock, agenda)')
b('3. El API Gateway')
b('4. El frontend')

# ════════════════════════════════════════════════════════════
# 5. ENDPOINTS
# ════════════════════════════════════════════════════════════
h1('5. Endpoints de la API')
p('Todas las rutas se acceden a través del API Gateway en http://localhost:8090. En la columna "Acceso requerido" se indica qué nivel de autenticación necesita cada endpoint.')

h2('5.1 Usuarios  (/gateway/users)')
endpoint_table([
    ('POST', '/gateway/users', 'Público', 'Registrar usuario nuevo'),
    ('GET', '/gateway/users/username/{username}', 'Público', 'Buscar usuario por username (login)'),
    ('GET', '/gateway/users', 'ADMIN', 'Listar todos los usuarios'),
    ('GET', '/gateway/users/{id}', 'ADMIN', 'Obtener usuario por ID'),
    ('GET', '/gateway/users/role/{role}', 'ADMIN', 'Filtrar usuarios por rol (USER / TECNICO / ADMIN)'),
    ('GET', '/gateway/users/active', 'ADMIN', 'Listar usuarios activos'),
    ('PUT', '/gateway/users/{id}', 'ADMIN', 'Actualizar datos del usuario (no modifica el rol)'),
    ('PATCH', '/gateway/users/{id}/assign-role', 'ADMIN', 'Asignar rol al usuario. Query param: ?role=TECNICO'),
    ('PATCH', '/gateway/users/{id}/toggle-status', 'ADMIN', 'Activar o desactivar cuenta de usuario'),
    ('DELETE', '/gateway/users/{id}', 'ADMIN', 'Eliminar usuario (el ADMIN no puede eliminarse)'),
])

h3('Cuerpo de la solicitud — Crear / actualizar usuario (JSON)')
p('{\n  "username": "jperez",\n  "email": "jperez@techfix.cl",\n  "password": "miClave123",\n  "fullName": "Juan Pérez",\n  "role": "USER"\n}')

h3('Valores válidos para el campo "role"')
table(['Valor', 'Descripción'],
    [('USER', 'Cliente estándar'), ('TECNICO', 'Personal técnico'), ('ADMIN', 'Administrador del sistema')])

h2('5.2 Pagos  (/gateway/pagos)')
endpoint_table([
    ('POST', '/gateway/pagos/nuevo/{idVisita}', 'ADMIN', 'Registrar pago para una visita técnica. Header opcional: Idempotency-Key'),
    ('GET', '/gateway/pagos/ver/{id}', 'ADMIN', 'Consultar datos de un pago por su ID'),
    ('PUT', '/gateway/pagos/alterar/{id}', 'ADMIN', 'Modificar un pago existente (no se puede modificar un pago ANULADO)'),
])

h3('Cuerpo de la solicitud — Crear / actualizar pago (JSON)')
p('{\n  "idVisitaTecnica": 12,\n  "monto": 45000.00,\n  "metodoPago": "TRANSFERENCIA",\n  "estadoPago": "PENDIENTE",\n  "descripcion": "Reparación pantalla notebook"\n}')

h3('Valores válidos para "metodoPago"')
table(['Valor', 'Descripción'],
    [('EFECTIVO','Pago en efectivo'),('TARJETA_DEBITO','Tarjeta de débito'),
     ('TARJETA_CREDITO','Tarjeta de crédito'),('TRANSFERENCIA','Transferencia bancaria')])

h3('Valores válidos para "estadoPago"')
table(['Valor', 'Descripción'],
    [('PENDIENTE','Cobro registrado pero no confirmado'),('PAGADO','Cobro confirmado'),('ANULADO','Cobro cancelado; no se puede modificar')])

h3('Protección contra cobros duplicados (Idempotencia)')
p('Al enviar una solicitud POST /nuevo/{id}, se puede incluir el header:')
p('  Idempotency-Key: 550e8400-e29b-41d4-a716-446655440000')
p('Si el servidor ya procesó esa clave antes, devuelve HTTP 200 con el pago original en lugar de crear uno nuevo. Si es un pago nuevo, responde HTTP 201. La clave debe ser un UUID v4 único por operación.')

h2('5.3 Inventario  (/gateway/stock)')

h3('Productos')
endpoint_table([
    ('GET', '/gateway/stock', 'Público', 'Listar productos paginados. Params: page, size, sortBy, direction'),
    ('GET', '/gateway/stock/{id}', 'Público', 'Obtener producto por ID'),
    ('GET', '/gateway/stock/sku/{sku}', 'Público', 'Obtener producto por código SKU'),
    ('GET', '/gateway/stock/search?q={texto}', 'Público', 'Buscar productos por texto. Params: page, size'),
    ('GET', '/gateway/stock/status/{status}', 'Público', 'Filtrar por estado (ACTIVE/INACTIVE/DISCONTINUED)'),
    ('GET', '/gateway/stock/category/{categoryId}', 'Público', 'Filtrar por categoría'),
    ('POST', '/gateway/stock', 'ADMIN', 'Crear producto nuevo'),
    ('PUT', '/gateway/stock/{id}', 'ADMIN', 'Actualizar producto'),
    ('DELETE', '/gateway/stock/{id}', 'ADMIN', 'Eliminar producto'),
])

h3('Movimientos de stock')
endpoint_table([
    ('POST', '/gateway/stock/movements', 'ADMIN', 'Registrar movimiento de inventario'),
    ('PATCH', '/gateway/stock/{id}/adjust-stock', 'ADMIN', 'Ajustar cantidad directamente'),
    ('GET', '/gateway/stock/{id}/movements', 'ADMIN', 'Ver historial de movimientos de un producto'),
])

h3('Alertas y estadísticas')
endpoint_table([
    ('GET', '/gateway/stock/alerts/low-stock', 'ADMIN', 'Listar productos con stock bajo el mínimo'),
    ('GET', '/gateway/stock/alerts/out-of-stock', 'ADMIN', 'Listar productos sin stock disponible'),
    ('GET', '/gateway/stock/stats', 'ADMIN', 'Ver estadísticas generales del inventario'),
])

h3('Categorías')
endpoint_table([
    ('GET', '/gateway/stock/categories', 'Público', 'Listar todas las categorías'),
    ('GET', '/gateway/stock/categories/{id}', 'Público', 'Obtener categoría por ID'),
    ('POST', '/gateway/stock/categories', 'ADMIN', 'Crear categoría'),
    ('PUT', '/gateway/stock/categories/{id}', 'ADMIN', 'Actualizar categoría'),
    ('DELETE', '/gateway/stock/categories/{id}', 'ADMIN', 'Eliminar categoría'),
])

h3('Cuerpo de la solicitud — Crear producto (JSON)')
p('{\n  "sku": "LAP-SCR-15-001",\n  "name": "Pantalla 15.6 FHD",\n  "description": "Pantalla de repuesto para notebook",\n  "brand": "LG",\n  "model": "LP156WFH",\n  "costPrice": 35000.00,\n  "salePrice": 55000.00,\n  "quantityInStock": 10,\n  "minStockLevel": 3,\n  "categoryId": 1,\n  "imageUrl": "https://ejemplo.com/img.jpg"\n}')

h3('Cuerpo de la solicitud — Registrar movimiento (JSON)')
p('{\n  "productId": 5,\n  "movementType": "PURCHASE",\n  "quantity": 20,\n  "unitCost": 35000.00,\n  "reference": "OC-2026-045",\n  "notes": "Compra a proveedor LG",\n  "createdBy": "admin"\n}')

h3('Valores válidos para "movementType"')
table(['Valor','Descripción'],
    [('PURCHASE','Compra / ingreso de stock'),('SALE','Venta / salida de stock'),
     ('ADJUSTMENT','Ajuste por inventario físico'),('RETURN','Devolución de cliente'),
     ('TRANSFER','Traslado entre bodegas'),('DAMAGE','Baja por producto dañado')])

h2('5.4 Citas  (/gateway/citas)')
endpoint_table([
    ('GET', '/gateway/citas', 'Autenticado', 'Listar todas las citas'),
    ('GET', '/gateway/citas/{id}', 'Autenticado', 'Obtener cita por ID'),
    ('GET', '/gateway/citas/cliente/{clienteId}', 'Autenticado', 'Citas de un cliente específico'),
    ('GET', '/gateway/citas/tecnico/{tecnicoId}', 'Autenticado', 'Citas asignadas a un técnico'),
    ('GET', '/gateway/citas/estado/{estado}', 'Autenticado', 'Filtrar citas por estado'),
    ('POST', '/gateway/citas', 'Autenticado', 'Crear cita nueva'),
    ('PATCH', '/gateway/citas/{id}/estado', 'ADMIN / TECNICO', 'Actualizar el estado de una cita'),
    ('DELETE', '/gateway/citas/{id}', 'ADMIN / TECNICO', 'Eliminar / cancelar cita'),
])

h3('Cuerpo de la solicitud — Crear cita (JSON)')
p('{\n  "fechaHora": "2026-06-15T10:00:00",\n  "tipoServicio": "REPARACION",\n  "descripcion": "Laptop no enciende",\n  "clienteId": 3,\n  "tecnicoId": 2\n}')

h3('Valores válidos para "tipoServicio"')
table(['Valor','Descripción'],[('REPARACION','Reparación de equipo'),('INSTALACION','Instalación de componente o software')])

h3('Valores válidos para "estado" (PATCH /estado)')
table(['Valor','Descripción'],
    [('PENDIENTE','Cita agendada sin atender'),('EN_PROCESO','Técnico en atención'),
     ('COMPLETADA','Visita finalizada'),('CANCELADA','Cita cancelada')])

h2('5.5 Clientes  (/gateway/clientes)')
endpoint_table([
    ('GET', '/gateway/clientes', 'Autenticado', 'Listar todos los clientes'),
    ('GET', '/gateway/clientes/{id}', 'Autenticado', 'Obtener cliente por ID'),
    ('POST', '/gateway/clientes', 'Autenticado', 'Registrar cliente nuevo'),
    ('PUT', '/gateway/clientes/{id}', 'Autenticado', 'Actualizar datos de un cliente'),
    ('DELETE', '/gateway/clientes/{id}', 'ADMIN', 'Eliminar cliente'),
])

h3('Cuerpo de la solicitud — Crear / actualizar cliente (JSON)')
p('{\n  "nombre": "María",\n  "apellido": "González",\n  "email": "mgonzalez@gmail.com",\n  "telefono": "+56912345678"\n}')

h2('5.6 Técnicos  (/gateway/tecnicos)')
endpoint_table([
    ('GET', '/gateway/tecnicos', 'Público', 'Listar todos los técnicos'),
    ('GET', '/gateway/tecnicos/activos', 'Público', 'Listar solo técnicos activos'),
    ('GET', '/gateway/tecnicos/{id}', 'Público', 'Obtener técnico por ID'),
    ('POST', '/gateway/tecnicos', 'ADMIN', 'Registrar técnico nuevo'),
    ('PUT', '/gateway/tecnicos/{id}', 'ADMIN', 'Actualizar datos de un técnico'),
    ('DELETE', '/gateway/tecnicos/{id}', 'ADMIN', 'Eliminar técnico'),
])

h3('Cuerpo de la solicitud — Crear / actualizar técnico (JSON)')
p('{\n  "nombre": "Carlos",\n  "apellido": "Reyes",\n  "email": "creyes@techfix.cl",\n  "telefono": "+56987654321",\n  "especialidad": "Electrónica y hardware",\n  "activo": true\n}')

# ════════════════════════════════════════════════════════════
# 6. CÓDIGOS DE RESPUESTA
# ════════════════════════════════════════════════════════════
h1('6. Códigos de respuesta HTTP')
table(
    ['Código', 'Significado', 'Cuándo ocurre'],
    [
        ('200 OK', 'Solicitud exitosa', 'Recurso encontrado o pago duplicado detectado (idempotencia)'),
        ('201 Created', 'Recurso creado', 'Se creó un nuevo usuario, pago, cita, etc.'),
        ('204 No Content', 'Eliminado exitosamente', 'DELETE exitoso'),
        ('400 Bad Request', 'Datos inválidos', 'Faltan campos obligatorios o formato incorrecto'),
        ('401 Unauthorized', 'No autenticado', 'Falta el header X-User-Id o el usuario no existe'),
        ('403 Forbidden', 'Sin permisos', 'El usuario existe pero no tiene el rol necesario'),
        ('404 Not Found', 'No encontrado', 'El ID consultado no existe en la base de datos'),
        ('409 Conflict', 'Conflicto', 'Email/username duplicado, ya existe un ADMIN, pago ANULADO'),
        ('503 Service Unavailable', 'Servicio caído', 'El servicio de autenticación no está disponible'),
    ]
)

# ════════════════════════════════════════════════════════════
# 7. SEGURIDAD
# ════════════════════════════════════════════════════════════
h1('7. Seguridad')
b('Autenticación por ID: el frontend envía el ID del usuario en el header X-User-Id de cada petición. El Gateway lo verifica en tiempo real.')
b('Control de acceso por rol: el Gateway bloquea solicitudes a recursos restringidos según el rol del usuario antes de que lleguen al servicio correspondiente.')
b('Idempotencia en pagos: evita cobros duplicados usando una clave única por operación, generada en el navegador.')
b('Bases de datos aisladas: cada módulo tiene su propia base de datos con credenciales independientes.')

# ════════════════════════════════════════════════════════════
# 8. ESTRUCTURA DEL REPOSITORIO
# ════════════════════════════════════════════════════════════
h1('8. Estructura del repositorio')
p('TechFix/\n'
  '├── docker-compose.yml\n'
  '├── .gitignore\n'
  '├── frontend/                  ← Código de la página web\n'
  '│   ├── Dockerfile\n'
  '│   ├── nginx.conf\n'
  '│   └── src/\n'
  '│       ├── pages/             ← Pantallas (Login, Home, Pagos, etc.)\n'
  '│       ├── components/        ← Navbar, Footer, tarjetas\n'
  '│       ├── hooks/             ← Lógica de conexión con los servidores\n'
  '│       ├── services/api.js    ← Configuración de las URLs del Gateway\n'
  '│       └── context/           ← Estado global de autenticación\n'
  '└── microservicios/\n'
  '    ├── gateway/\n'
  '    ├── roles/\n'
  '    ├── pagos/\n'
  '    ├── stock/\n'
  '    └── agenda/')

out = '/home/enrique/Documentos/GitHub/TechFix/TechFix_Documentacion.docx'
doc.save(out)
print(f'Documento guardado en: {out}')
